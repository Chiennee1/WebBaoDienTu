# -*- coding: utf-8 -*-
import re
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement

MD = Path(r"e:\Project_Buy\LapTrinhWebsite\WebBaoDienTu\baoDienTu\Docs\PHAN_CONG_NHIEM_VU_VA_KE_HOACH_TRIEN_KHAI.md")
OUT = Path(r"e:\Project_Buy\LapTrinhWebsite\WebBaoDienTu\baoDienTu\Docs\PHAN_CONG_NHIEM_VU_VA_KE_HOACH_TRIEN_KHAI.docx")


def set_cell_shading(cell, fill_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, bold=False, italic=False, size=11, mono=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Consolas" if mono else "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_formatted_text(paragraph, text, base_size=11):
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=base_size)
        chunk = match.group(0)
        if chunk.startswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            set_run_font(run, bold=True, size=base_size)
        elif chunk.startswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            set_run_font(run, mono=True, size=base_size - 1)
        else:
            run = paragraph.add_run(chunk[1:-1])
            set_run_font(run, italic=True, size=base_size)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size)


def parse_table_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_table_separator(line):
    return bool(re.match(r"^\|\s*[-: ]+\|\s*[-: |]+\s*$", line.strip()))


def add_table(doc, rows):
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(col_count):
            cell = table.rows[i].cells[j]
            text = row[j] if j < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            add_formatted_text(p, text, base_size=10)
            if i == 0:
                for run in p.runs:
                    run.bold = True
                set_cell_shading(cell, "D9E2F3")
    doc.add_paragraph()


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("\n".join(lines))
    set_run_font(run, mono=True, size=9)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F2F2")
    shading.set(qn("w:val"), "clear")
    p._p.get_or_add_pPr().append(shading)


def convert():
    lines = MD.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    i = 0
    table_buffer = []
    code_buffer = []
    in_code = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_buffer)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue
        if stripped.startswith("|") and "|" in stripped[1:]:
            if not is_table_separator(stripped):
                table_buffer.append(parse_table_row(stripped))
            i += 1
            continue
        elif table_buffer:
            add_table(doc, table_buffer)
            table_buffer = []
        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue
        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:].strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75)
            add_formatted_text(p, stripped[2:].strip(), base_size=12)
            for run in p.runs:
                run.italic = True
            i += 1
            continue
        if stripped.startswith("- [ ] "):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_text(p, "☐ " + stripped[6:].strip(), base_size=12)
            i += 1
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_text(p, stripped[2:].strip(), base_size=12)
            i += 1
            continue
        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_formatted_text(p, re.sub(r"^\d+\.\s", "", stripped), base_size=12)
            i += 1
            continue
        if stripped == "":
            i += 1
            continue
        p = doc.add_paragraph()
        add_formatted_text(p, stripped, base_size=13)
        i += 1

    if table_buffer:
        add_table(doc, table_buffer)
    if code_buffer:
        add_code_block(doc, code_buffer)

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(
        "Tài liệu phân công — Dự án WebBaoDienTu — Nhóm 5 thành viên\n"
        "Trường Đại học Mở Hà Nội — Khoa Công nghệ Thông tin — Môn Lập trình Web"
    )
    set_run_font(run, italic=True, size=11)
    doc.save(str(OUT))
    print(f"OK {OUT} size={OUT.stat().st_size}")


if __name__ == "__main__":
    convert()
